"""
Document processing service.
Handles extraction of text from PDF and DOCX documents.
"""

from __future__ import annotations

import logging
import re
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Dict, List
from zipfile import ZipFile


logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Main document processing class."""

    def __init__(self) -> None:
        self.supported_formats = [".pdf", ".docx"]

    def process(self, file_path: str) -> Dict:
        """
        Process any supported document.

        Args:
            file_path: Path to the document.

        Returns:
            Dictionary with extracted text and metadata.
        """
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {file_ext}. Supported: {self.supported_formats}")

        if file_ext == ".pdf":
            return self._process_pdf(file_path)
        return self._process_docx(file_path)

    def _process_pdf(self, file_path: str) -> Dict:
        """Extract text from a PDF document."""
        logger.info("Processing PDF: %s", file_path)

        try:
            from PyPDF2 import PdfReader
        except ModuleNotFoundError:
            message = "PyPDF2 is not installed. Install dependencies from requirements.txt."
            logger.error(message)
            return {"success": False, "error": message, "file_type": "pdf"}

        try:
            reader = PdfReader(file_path)
            pages: List[Dict] = []
            text_chunks: List[str] = []

            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                pages.append(
                    {
                        "page_num": page_num,
                        "text": page_text,
                        "char_count": len(page_text),
                    }
                )
                if page_text:
                    text_chunks.append(page_text)

            metadata = reader.metadata or {}
            return {
                "success": True,
                "file_type": "pdf",
                "pages": len(pages),
                "text": "\n".join(text_chunks),
                "pages_detail": pages,
                "metadata": {
                    "author": self._safe_metadata_value(metadata.get("/Author")),
                    "creator": self._safe_metadata_value(metadata.get("/Creator")),
                    "producer": self._safe_metadata_value(metadata.get("/Producer")),
                },
            }
        except Exception as exc:  # noqa: BLE001
            logger.exception("PDF processing failed")
            return {"success": False, "error": str(exc), "file_type": "pdf"}

    def _process_docx(self, file_path: str) -> Dict:
        """Extract text from a DOCX document."""
        logger.info("Processing DOCX: %s", file_path)

        try:
            import docx
        except ModuleNotFoundError:
            logger.warning("python-docx not installed. Falling back to stdlib parser.")
            return self._process_docx_with_stdlib(file_path)

        try:
            doc = docx.Document(file_path)
            paragraphs: List[Dict] = []
            full_text_parts: List[str] = []

            for para_num, para in enumerate(doc.paragraphs, start=1):
                text_value = para.text.strip()
                if not text_value:
                    continue

                paragraphs.append(
                    {
                        "para_num": para_num,
                        "text": text_value,
                        "style": para.style.name if para.style else "Normal",
                    }
                )
                full_text_parts.append(text_value)

            tables: List[Dict] = []
            for table_num, table in enumerate(doc.tables, start=1):
                table_data: List[List[str]] = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    if any(row_data):
                        table_data.append(row_data)

                if table_data:
                    tables.append({"table_num": table_num, "data": table_data})

            core_properties = {}
            for property_name in (
                "author",
                "category",
                "comments",
                "content_status",
                "created",
                "identifier",
                "keywords",
                "language",
                "last_modified_by",
                "last_printed",
                "modified",
                "revision",
                "subject",
                "title",
                "version",
            ):
                value = getattr(doc.core_properties, property_name, None)
                if value not in (None, ""):
                    core_properties[property_name] = str(value)

            return {
                "success": True,
                "file_type": "docx",
                "paragraphs": len(paragraphs),
                "text": "\n".join(full_text_parts),
                "paragraphs_detail": paragraphs,
                "tables": tables,
                "metadata": {
                    "parser": "python-docx",
                    "core_properties": core_properties,
                },
            }
        except Exception as exc:  # noqa: BLE001
            logger.exception("DOCX processing failed")
            return {"success": False, "error": str(exc), "file_type": "docx"}

    def _process_docx_with_stdlib(self, file_path: str) -> Dict:
        """
        Fallback DOCX parser using stdlib only.
        Works when python-docx is not installed.
        """
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        try:
            with ZipFile(file_path) as archive:
                document_xml = archive.read("word/document.xml")
        except Exception as exc:  # noqa: BLE001
            return {"success": False, "error": f"Failed to read DOCX XML: {exc}", "file_type": "docx"}

        try:
            root = ET.fromstring(document_xml)
            body = root.find("w:body", namespace)
            paragraph_nodes = body.findall("w:p", namespace) if body is not None else []
            table_nodes = body.findall("w:tbl", namespace) if body is not None else []

            paragraphs: List[Dict] = []
            text_parts: List[str] = []
            for para_num, para in enumerate(paragraph_nodes, start=1):
                para_text = self._extract_xml_text(para, namespace)
                if not para_text:
                    continue
                paragraphs.append({"para_num": para_num, "text": para_text, "style": "Unknown"})
                text_parts.append(para_text)

            tables: List[Dict] = []
            for table_num, table in enumerate(table_nodes, start=1):
                table_data: List[List[str]] = []
                for row in table.findall("w:tr", namespace):
                    row_data: List[str] = []
                    for cell in row.findall("w:tc", namespace):
                        cell_parts: List[str] = []
                        for para in cell.findall(".//w:p", namespace):
                            para_text = self._extract_xml_text(para, namespace)
                            if para_text:
                                cell_parts.append(para_text)
                        row_data.append(" ".join(cell_parts).strip())

                    if any(row_data):
                        table_data.append(row_data)

                if table_data:
                    tables.append({"table_num": table_num, "data": table_data})

            return {
                "success": True,
                "file_type": "docx",
                "paragraphs": len(paragraphs),
                "text": "\n".join(text_parts),
                "paragraphs_detail": paragraphs,
                "tables": tables,
                "metadata": {
                    "parser": "xml-fallback",
                    "core_properties": {},
                },
            }
        except Exception as exc:  # noqa: BLE001
            logger.exception("Fallback DOCX parsing failed")
            return {"success": False, "error": str(exc), "file_type": "docx"}

    def clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""

        text = text.replace("\r\n", "\n").replace("\r", "\n")
        replacements = {
            "\u00ef\u00bf\u00bd": "'",
            "\u00ef\u00ac\u20ac": "ff",
            "\u00ef\u00ac\u0081": "fi",
            "\u00ef\u00ac\u201a": "fl",
            "\u00c2\u00bd": "1/2",
            "\u00c2\u00bc": "1/4",
            "\u00c2\u00be": "3/4",
            "\u00e2\u20ac\u00a2": "-",
        }
        for old, new in replacements.items():
            text = text.replace(old, new)

        text = re.sub(r"(?im)^\s*page\s+\d+\s+of\s+\d+\s*$", "", text)
        text = re.sub(r"(?m)^\s*\d+\s*$", "", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"[ \t]*\n[ \t]*", "\n", text)
        text = re.sub(r"(?<![.!?:;\n])\n(?!\n)", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def detect_questions(self, text: str) -> List[Dict]:
        """
        Detect question boundaries in ACCA exam answers.

        Looks for patterns like:
        - Requirement (a) - 4 marks
        - Question 1
        - (a) (4 marks)
        """
        questions: List[Dict] = []

        requirement_pattern = re.compile(
            r"(?:Requirement|Req)[\s.:]*\(?([a-z])\)?\s*(?:[-:]\s*)?(\d+)\s*marks?",
            re.IGNORECASE,
        )
        question_pattern = re.compile(r"Question\s+(\d+)", re.IGNORECASE)
        part_pattern = re.compile(r"\(([a-z])\)\s*\((\d+)\s*marks?\)", re.IGNORECASE)

        lines = text.split("\n")
        current_question: Dict | None = None
        question_text: List[str] = []

        for line_number, line in enumerate(lines):
            match_requirement = requirement_pattern.search(line)
            match_question = question_pattern.search(line)
            match_part = part_pattern.search(line)

            if match_requirement or match_question or match_part:
                if current_question is not None:
                    current_question["text"] = "\n".join(question_text).strip()
                    questions.append(current_question)

                if match_requirement:
                    current_question = {
                        "type": "requirement",
                        "part": match_requirement.group(1).lower(),
                        "marks": int(match_requirement.group(2)),
                        "start_line": line_number,
                        "header": line.strip(),
                    }
                elif match_question:
                    current_question = {
                        "type": "question",
                        "number": int(match_question.group(1)),
                        "start_line": line_number,
                        "header": line.strip(),
                    }
                else:
                    current_question = {
                        "type": "part",
                        "part": match_part.group(1).lower(),
                        "marks": int(match_part.group(2)),
                        "start_line": line_number,
                        "header": line.strip(),
                    }
                question_text = []
                continue

            if current_question is not None:
                question_text.append(line)

        if current_question is not None:
            current_question["text"] = "\n".join(question_text).strip()
            questions.append(current_question)

        return questions

    def extract_marking_points(self, text: str) -> List[Dict]:
        """
        Extract individual marking points from marking scheme text.

        Looks for:
        - 1 mark for X
        - 1/2 mark for Y
        - Award 1 mark if...
        """
        points: List[Dict] = []
        seen_keys = set()

        patterns = [
            re.compile(r"(?<![/\d])(\d+(?:\.\d+)?)\s*mark(?:s)?\s*(?:for|if|when)\s*([^.\n]+)", re.IGNORECASE),
            re.compile(r"(1/2|1/4|3/4|\u00bd|\u00bc|\u00be|\u00c2\u00bd|\u00c2\u00bc|\u00c2\u00be)\s*mark(?:s)?\s*(?:for|if|when)\s*([^.\n]+)", re.IGNORECASE),
            re.compile(r"Award\s*(\d+(?:\.\d+)?)\s*mark(?:s)?\s*(?:for|if|when)\s*([^.\n]+)", re.IGNORECASE),
        ]

        for pattern in patterns:
            for mark_text, criteria in pattern.findall(text):
                normalized_criteria = criteria.strip()
                key = (self._parse_marks(mark_text), normalized_criteria.lower())
                if key in seen_keys:
                    continue
                seen_keys.add(key)
                points.append(
                    {
                        "marks": self._parse_marks(mark_text),
                        "criteria": normalized_criteria,
                        "source": "explicit",
                    }
                )

        bullet_pattern = re.compile(r"[\u2022\-*]\s*([^.\n]+)")
        for bullet in bullet_pattern.findall(text):
            criteria = bullet.strip()
            if len(criteria.split()) < 4:
                continue
            key = (1.0, criteria.lower())
            if key in seen_keys:
                continue
            seen_keys.add(key)
            points.append(
                {
                    "marks": 1.0,
                    "criteria": criteria,
                    "source": "bullet",
                    "confidence": "low",
                }
            )

        return points

    def _parse_marks(self, mark_str: str) -> float:
        """Convert mark strings to float values."""
        normalized = mark_str.strip().lower().replace(" ", "")
        fractions = {
            "1/2": 0.5,
            "1/4": 0.25,
            "3/4": 0.75,
            "\u00bd": 0.5,
            "\u00bc": 0.25,
            "\u00be": 0.75,
            "\u00c2\u00bd": 0.5,
            "\u00c2\u00bc": 0.25,
            "\u00c2\u00be": 0.75,
        }
        if normalized in fractions:
            return fractions[normalized]

        try:
            return float(normalized)
        except ValueError:
            return 1.0

    def extract_professional_marks(self, text: str) -> Dict:
        """
        Extract professional marks criteria.

        Professional marks are usually for:
        - Presentation
        - Structure
        - Clarity
        - Professional language
        """
        criteria = {
            "presentation": 0,
            "structure": 0,
            "clarity": 0,
            "professional_language": 0,
            "total": 0,
        }

        prof_pattern = re.compile(r"(?:professional|presentation|structure|clarity)[^.]*?(\d+)\s*marks?", re.IGNORECASE)
        matches = prof_pattern.findall(text)
        if matches:
            criteria["total"] = sum(int(mark) for mark in matches)
        else:
            criteria["total"] = 4
        return criteria

    def _extract_xml_text(self, node: ET.Element, namespace: Dict[str, str]) -> str:
        """Extract all text values from an XML node."""
        text_parts: List[str] = []
        for text_node in node.findall(".//w:t", namespace):
            if text_node.text:
                text_parts.append(text_node.text)
        return "".join(text_parts).strip()

    def _safe_metadata_value(self, value: object) -> str:
        """Normalize metadata values from PDF readers."""
        if value in (None, ""):
            return "Unknown"
        return str(value)

